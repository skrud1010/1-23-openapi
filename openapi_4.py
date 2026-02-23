import streamlit as st
import os
import requests
import zipfile
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime, timedelta
import streamlit as st
from openai import OpenAI


# 1. 환경 설정
load_dotenv()
EXCHANGE_KEY = os.getenv("EXCHANGE_RATE_KEY")
OPENAI_KEY = st.secrets["OPENAI_KEY"]
client = OpenAI(api_key=OPENAI_KEY)

# --- [서류 생성 엔진] 각 서류 양식에 맞게 표 구조 변경 ---
def create_custom_trade_doc(data, doc_type):
    doc = Document()
    title_text = doc_type.replace('_', ' ').upper()
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 상단 기본 정보 표
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = f"SHIPPER/SELLER:\n{data['exp_name']}\n{data['exp_addr']}\n\nCONSIGNEE:\n{data['imp_name']}\n{data['imp_addr']}"
    
    # 서류별 상단 우측 정보 (인보이스 번호, L/C 번호 등 반영) 
    right_info = f"REF/INV NO: {data['inv_no']}\nDATE: {data['date']}\n"
    if doc_type in ["Commercial_Invoice", "Packing_List"]:
        right_info += f"L/C NO: {data['contract_no']}\n"
    right_info += f"PAYMENT: {data['payment_terms']}"
    table.rows[0].cells[1].text = right_info

    doc.add_paragraph("\n")

    # --- 서류별 특화 항목 구성 ---
    if doc_type == "Packing_List":
        # 포장명세서: 수량, 순중량, 총중량, 용적 항목 포함 
        item_table = doc.add_table(rows=1, cols=5)
        item_table.style = 'Table Grid'
        hdrs = item_table.rows[0].cells
        hdrs[0].text, hdrs[1].text = 'DESCRIPTION', 'QTY'
        hdrs[2].text, hdrs[3].text, hdrs[4].text = 'NET WT', 'GROSS WT', 'MEASUREMENT'
        
        row = item_table.add_row().cells
        row[0].text = data['product_name']
        row[1].text = f"{data['qty']:,}"
        row[2].text = "1,208.06 KGS" # 예시 데이터 
        row[3].text = "1,317 KGS"    # 예시 데이터 
        row[4].text = "24.5 CBM"     # 예시 데이터 

    elif doc_type == "Bill_of_Lading":
        # 선하증권: 선박명, 선적항, 하역항 정보 강조
        doc.add_paragraph(f"VESSEL/FLIGHT: {data['vessel']}\nFROM: {data['pol']}\nTO: {data['pod']}").bold = True
        item_table = doc.add_table(rows=1, cols=3)
        item_table.style = 'Table Grid'
        hdrs = item_table.rows[0].cells
        hdrs[0].text, hdrs[1].text, hdrs[2].text = 'DESCRIPTION', 'MARKS & NOS', 'TOTAL WT'
        
        row = item_table.add_row().cells
        row[0].text = data['product_name']
        row[1].text = f"C/NO.1-{data['qty']}"
        row[2].text = "1,317 KGS"

    else:
        # 기타(S/C, C/I, L/C): 제품명, 수량, 단가, 총액 표 
        item_table = doc.add_table(rows=1, cols=4)
        item_table.style = 'Table Grid'
        hdrs = item_table.rows[0].cells
        hdrs[0].text, hdrs[1].text = 'DESCRIPTION', 'QTY'
        hdrs[2].text, hdrs[3].text = 'UNIT PRICE', 'AMOUNT'
        
        row = item_table.add_row().cells
        row[0].text = data['product_name']
        row[1].text = f"{data['qty']:,}"
        row[2].text = f"{data['currency']} {data['unit_p']:,.2f}"
        row[3].text = f"{data['currency']} {data['total_f']:,.2f}"

    doc.add_paragraph("\n")
    # 하단 뱅킹 및 서명 정보
    doc.add_paragraph(f"BANK DETAILS:\n{data['bank_details']}").style.font.size = Pt(9)
    doc.add_paragraph("\nAuthorized Signature: ________________________")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- Streamlit 실행 부분 ---
st.set_page_config(page_title="통합 무역 서류 생성기", layout="wide")
st.title("📂 5대 무역 서류(Set) 맞춤 양식 생성기")

with st.form("trade_form"):
    # ... (기존과 동일한 입력 필드 유지) ...
    col1, col2 = st.columns(2)
    with col1:
        exp_name = st.text_input("공급사명", value="GILDING TRADING CO., LTD.") # 예시 
        imp_name = st.text_input("구매사명", value="MONARCH PRODUCTS CO., LTD.") # 예시 
        inv_no = st.text_input("참조 번호", value="8905 BK 1007")
    with col2:
        vessel = st.text_input("선박명", value="PHOENIX") # 예시 
        pol = st.text_input("선적항", value="BUSAN, KOREA") # 예시 
        pod = st.text_input("하역항", value="DETROIT, U.S.A") # 예시 
    
    submitted = st.form_submit_button("📦 5종 서류 일괄 생성 (맞춤 양식)")

if submitted:
    # ... (환율 계산 및 데이터 정리 로직 동일) ...
    with st.spinner("서류 생성 중..."):
        data = {
            "exp_name": exp_name, "exp_addr": "159, SAMSUNG-DONG, SEOUL", 
            "imp_name": imp_name, "imp_addr": "DETROIT, MICHIGAN 48203",
            "inv_no": inv_no, "contract_no": "L/C 55352", "date": "MAY. 20, 2007",
            "vessel": vessel, "pol": pol, "pod": pod, "currency": "USD",
            "payment_terms": "L/C AT SIGHT", "product_name": "NYLON OXFORD",
            "qty": 60000, "unit_p": 1.00, "total_f": 60000.00,
            "bank_details": "Bank: Korea Bank\nSwift: KORBKRSE", "incoterms": "FOB BUSAN"
        }

        file_types = ["Bill_of_Lading", "Commercial_Invoice", "Letter_of_Credit", "Packing_List", "Sales_Contract"]
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for f_type in file_types:
                doc_stream = create_custom_trade_doc(data, f_type)
                zf.writestr(f"{f_type}_{inv_no}.docx", doc_stream.getvalue())
        
        st.success("✅ 각 양식에 맞춘 5종 서류 세트가 준비되었습니다!")
        st.download_button("📥 모든 서류 ZIP 다운로드", zip_buffer.getvalue(), f"Trade_Docs_{inv_no}.zip")